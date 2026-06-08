#!/usr/bin/env python3
"""Convierte Markdown básico a .docx (títulos, tablas, listas, negrita)."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_rich_paragraph(doc, text, style=None, size=11):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, size=size - 1, name="Consolas")
        else:
            run = p.add_run(part)
            set_run_font(run, size=size)
    return p


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_sep(line):
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", line.strip()))


def md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    i = 0
    in_code = False
    code_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                set_run_font(run, size=9, name="Consolas")
                p.paragraph_format.left_indent = Inches(0.25)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            title = stripped.lstrip("#").strip()
            if level == 1:
                p = doc.add_heading(title, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=18, bold=True, color=RGBColor(0x9A, 0x16, 0x20))
            elif level == 2:
                doc.add_heading(title, level=1)
            elif level == 3:
                doc.add_heading(title, level=2)
            else:
                doc.add_heading(title, level=3)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            cols = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=cols)
            table.style = "Table Grid"
            for c, h in enumerate(headers):
                cell = table.rows[0].cells[c]
                cell.text = ""
                run = cell.paragraphs[0].add_run(h)
                set_run_font(run, size=10, bold=True)
            for r, row in enumerate(rows):
                for c in range(cols):
                    val = row[c] if c < len(row) else ""
                    val = re.sub(r"\*\*([^*]+)\*\*", r"\1", val)
                    val = val.replace("`", "")
                    table.rows[r + 1].cells[c].text = val
            doc.add_paragraph()
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            add_rich_paragraph(doc, text, style="List Bullet")
            i += 1
            continue

        m_num = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if m_num:
            add_rich_paragraph(doc, m_num.group(2), style="List Number")
            i += 1
            continue

        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            p = add_rich_paragraph(doc, text)
            p.paragraph_format.left_indent = Inches(0.3)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        add_rich_paragraph(doc, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"OK: {docx_path}")


def main():
    root = Path(__file__).resolve().parents[1]
    targets = [
        ("docs/INSTRUCTIVO_USO_APLICACION.md", "docs/INSTRUCTIVO_USO_APLICACION.docx"),
        ("docs/DOCUMENTACION_SISTEMA.md", "docs/DOCUMENTACION_SISTEMA.docx"),
    ]
    if len(sys.argv) > 1:
        for arg in sys.argv[1:]:
            p = Path(arg)
            md_to_docx(p, p.with_suffix(".docx"))
        return
    for md_rel, docx_rel in targets:
        md = root / md_rel
        if md.exists():
            md_to_docx(md, root / docx_rel)


if __name__ == "__main__":
    main()
